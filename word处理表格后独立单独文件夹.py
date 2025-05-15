import os
import re
import pandas as pd
from docx import Document
import shutil


def sanitize_filename(filename):
    """清理文件名中的非法字符"""
    # 保留字母、数字、中文、下划线、空格
    return re.sub(r'[\\/*?:"<>|]', '', filename).replace(' ', '_')


def process_word_files_in_folder(input_folder, output_base_path):
    # 确保输出基础路径存在
    os.makedirs(output_base_path, exist_ok=True)
    print(f"正在处理文件夹：{input_folder}")
    print(f"输出文件夹：{output_base_path}")

    # 获取文件夹内所有Word文件
    word_files = [f for f in os.listdir(input_folder) if f.endswith('.docx')]
    print(f"找到的Word文件：{word_files}")
    
    if not word_files:
        print(f"错误：在文件夹 {input_folder} 中没有找到 .docx 文件！")
        return

    for word_file in word_files:
        word_path = os.path.join(input_folder, word_file)
        print(f"\n正在处理：{word_file}")

        try:
            # 创建与Word文件同名的输出文件夹
            base_name = os.path.splitext(word_file)[0]
            output_folder = os.path.join(output_base_path, base_name)
            os.makedirs(output_folder, exist_ok=True)

            # 复制原始Word文件到输出文件夹
            shutil.copy2(word_path, output_folder)

            # 读取Word网页中的所有表格
            doc = Document(word_path)
            all_table_data = []
            print(f"文件中的表格数量：{len(doc.tables)}")

            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    table_data.append([cell.text.strip() for cell in row.cells])
                all_table_data.append(table_data)

            # 生成与Word文件同名的Excel文件
            excel_file_name = f"{base_name}.xlsx"
            excel_file_path = os.path.join(output_folder, excel_file_name)

            with pd.ExcelWriter(excel_file_path, engine="openpyxl") as writer:
                for i, table in enumerate(all_table_data):
                    # 获取表格Sheet名（优先使用A1单元格内容）
                    if table and table[0] and table[0][0].strip():
                        sheet_name = sanitize_filename(table[0][0].strip())
                    else:
                        sheet_name = f"Sheet_{i + 1}"

                    # 处理重复Sheet名
                    unique_sheet_name = sheet_name
                    suffix = 1
                    while unique_sheet_name in writer.sheets:
                        unique_sheet_name = f"{sheet_name}_{suffix}"
                        suffix += 1

                    print(f"正在处理第 {i+1} 个表格，Sheet名称：{unique_sheet_name}")
                    df = pd.DataFrame(table[1:], columns=table[0])
                    df.to_excel(writer, sheet_name=unique_sheet_name, index=False)

            print(f"完成处理：{word_file}")
            print(f"Excel文件已保存到：{excel_file_path}")

        except Exception as e:
            print(f"处理文件 {word_file} 时出错：{str(e)}")


if __name__ == "__main__":
    print("程序开始执行...")
    # 获取当前脚本所在目录的绝对路径
    current_dir = os.path.dirname(os.path.abspath(__file__))
    print(f"当前工作目录：{current_dir}")
    
    # 处理文件
    process_word_files_in_folder(current_dir, os.path.join(current_dir, '处理结果'))
    print("\n程序执行完成！")
    print("请检查'处理结果'文件夹查看输出文件。")
    input("按回车键退出...")