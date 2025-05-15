# 批量提取word中的表格，放到同一个文件夹即可批量处理，表格名字为word对应名字，表中sheet为表对应名字，可以按需提取
#注意事项：word和py文件放到同一个文件夹
# author: yuchenqiang


import os
import re
import pandas as pd
from docx import Document


def process_word_files_in_folder(folder_path):
    print(f"开始处理文件夹：{folder_path}")
    # 获取文件夹内所有Word文件
    word_files = [f for f in os.listdir(folder_path) if f.endswith('.docx')]
    print(f"找到的Word文件：{word_files}")

    for word_file in word_files:
        print(f"\n开始处理文件：{word_file}")
        # 构建Word文件完整路径
        word_path = os.path.join(folder_path, word_file)
        print(f"文件完整路径：{word_path}")

        try:
            # 读取Word网页中的所有表格
            doc = Document(word_path)
            all_table_data = []
            print(f"文件中的表格数量：{len(doc.tables)}")

            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    table_data.append([cell.text.strip() for cell in row.cells])
                all_table_data.append(table_data)

            # 生成Excel文件名（与Word文件名一致）
            excel_file_name = re.sub(r'[\\/*?:"<>|]', '', os.path.splitext(word_file)[0]) + '.xlsx'
            excel_file_path = os.path.join(folder_path, excel_file_name)
            print(f"将生成Excel文件：{excel_file_name}")

            # 创建Excel写入对象
            with pd.ExcelWriter(excel_file_path, engine="openpyxl") as writer:
                # 处理每个表格并生成对应Sheet
                for i, table in enumerate(all_table_data):
                    # 获取表格的Sheet名（优先使用A1单元格内容）
                    if table and table[0] and table[0][0].strip():
                        sheet_name = table[0][0].strip()
                    else:
                        sheet_name = f'Table_{i + 1}'

                    # 处理特殊字符
                    sheet_name = re.sub(r'[\\/*?:"<>|]', '', sheet_name).replace(' ', '_')
                    print(f"处理表格 {i+1}，Sheet名称：{sheet_name}")

                    # 将表格数据转换为DataFrame并写入Excel
                    df = pd.DataFrame(table[1:], columns=table[0])
                    df.to_excel(writer, sheet_name=sheet_name, index=False)

            print(f"已处理文件：{word_file}，生成Excel：{excel_file_name}")
        except Exception as e:
            print(f"处理文件 {word_file} 时出错：{str(e)}")


if __name__ == "__main__":
    print("程序开始执行...")
    # 获取脚本所在目录的绝对路径
    script_dir = os.path.dirname(os.path.abspath(__file__))
    print(f"脚本所在目录：{script_dir}")
    # 使用脚本所在目录作为工作目录
    process_word_files_in_folder(script_dir)
    print("程序执行完成！")