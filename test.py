# import docx
# import pandas as pd
#
# doc = docx.Document("202502-氨氯地平贝那普利产品上市后研究报告.docx")
# data = []
#
# for table in doc.tables:
#     table_data = []
#     for row in table.rows:
#         table_data.append([cell.text for cell in row.cells])
#     data.append(table_data)
#
# # 将每个表格保存为单独的Sheet
# with pd.ExcelWriter("output.xlsx", engine="openpyxl") as writer:
#     for i, table in enumerate(data):
#         df = pd.DataFrame(table[1:], columns=table[0])  # 假设第一行为表头
#         df.to_excel(writer, sheet_name=f'Sheet{i+1}', index=False)

import docx
import pandas as pd

import docx
import pandas as pd
import re
doc = docx.Document("202502-氨氯地平贝那普利产品上市后研究报告.docx")
data = []

for table in doc.tables:
    table_data = []
    for row in table.rows:
        table_data.append([cell.text for cell in row.cells])
    data.append(table_data)

# 将每个表格保存为单独的Sheet，Sheet名取自对应表的A1单元格
with pd.ExcelWriter("output.xlsx", engine="openpyxl") as writer:
    for i, table in enumerate(data):
        # 获取A1单元格的值作为Sheet名（假设A1在表头第一列）
        if table and table[0] and table[0][0].strip():  # 检查A1是否存在且非空
            sheet_name = table[0][0].strip()  # 去除首尾空格
        else:
            sheet_name = f'Sheet{i + 1}'  # 默认Sheet名

        # 处理特殊字符（如空格、斜杠等）
        sheet_name = re.sub(r'[\\/*?:"<>|]', '', sheet_name)
        sheet_name = sheet_name.replace(' ', '_')

        df = pd.DataFrame(table[1:], columns=table[0])
        df.to_excel(writer, sheet_name=sheet_name, index=False)